import streamlit as st
from openai import OpenAI
from apikey import *
import os
import shutil
import requests
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image
import torch
import random
import time
import cv2
import docx2pdf
import scipy
import numpy as np
from pydub import AudioSegment
from moviepy.editor import VideoFileClip, AudioFileClip, concatenate_videoclips, concatenate_audioclips, CompositeAudioClip
from audiocraft.models import musicgen
from audiocraft.utils.notebook import display_audio
import torch
from scipy.io import wavfile

from utils import *
if os.name == 'nt':
    import pythoncom
    pythoncom.CoInitialize()
if os.name == 'posix':
    import subprocess
    def convert_docx_to_pdf(docx_path, pdf_path):
        subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path])

client = OpenAI(api_key=decrypt(OPENAI_API_KEY))

def get_image_data(url):
    response = requests.get(url)
    response.raise_for_status()
    return response.content

if 'chapter_count' not in st.session_state:
    st.session_state.chapter_count = 1
if 'file_available' not in st.session_state:
    st.session_state.file_available = False
if 'images' not in st.session_state:
    st.session_state.images = []
if 'cur_time' not in st.session_state:
    st.session_state.cur_time = 0
if 'screenplays' not in st.session_state:
    st.session_state.screenplays = []
if 'subtitles' not in st.session_state:
    st.session_state.subtitles = []
# initialize document format
screenplay_document = Document()
section = screenplay_document.sections[0]
section.page_width = Inches(8.27)  # Width for A4 paper
section.page_height = Inches(11.69)  # Height for A4 paper

section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
screenplay_document = Document()

subtitle_document = Document()
section = subtitle_document.sections[0]
section.page_width = Inches(8.27)  # Width for A4 paper
section.page_height = Inches(11.69)  # Height for A4 paper

section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)


st.write('Film Title')
content_text = st.text_input('Write your contents of your film')
chapter_count = st.text_input('Write the count of parts')
col1, col2, col3, col4, col5 = st.columns(5)
man_count = col2.text_input('Write the count of man characters')
woman_count = col4.text_input('Write the count of woman characters')
generate_btn = st.button('Generate screenplay and subtitles')

if content_text and generate_btn and chapter_count and man_count and woman_count:
    st.session_state.file_available = False
    st.session_state.chapter_count = int(chapter_count)
    st.session_state.cur_time = random.randint(1000, 3000)
    st.session_state.screenplays = []
    st.session_state.subtitles = []
    man_count = int(man_count)
    woman_count = int(woman_count)
    character_count = man_count + woman_count
    man, woman = characterNames(man_count, woman_count)
    character = man + woman
    # for i in range(man_count):
    #     start_time = time.time()
    #     response = client.images.generate(
    #         model="dall-e-3",
    #         prompt=f'one cartoon man character, not repeat, white background, no background, whole body, color image',
    #         size="1024x1792",
    #         quality="standard",
    #         n=1,
    #         )
    #     image_data = requests.get(response.data[0].url).content
    #     image_data = Image.open(BytesIO(image_data))
    #     image_data.save('out.png')
    #     character = removeBackground('out.png')
    #     character = cv2.resize(character, (character.shape[1] // (character_count + 1), character.shape[0] // (character_count + 1)), interpolation = cv2.INTER_AREA)
    #     cv2.imwrite(f'character/{i + 1}.png', character)
    #     os.remove('out.png')
    #     wait_time = time.time() - start_time
    #     if wait_time < 60:
    #         time.sleep(60 - wait_time)

    
    # for i in range(woman_count):
    #     start_time = time.time()
    #     response = client.images.generate(
    #         model="dall-e-3",
    #         prompt=f'one cartoon woman character, not repeat, white background, no background, whole body, color image',
    #         size="1024x1792",
    #         quality="standard",
    #         n=1,
    #         )
    #     image_data = requests.get(response.data[0].url).content
    #     image_data = Image.open(BytesIO(image_data))
    #     image_data.save('out.png')
    #     character = removeBackground('out.png')
    #     character = cv2.resize(character, (character.shape[1] // (character_count + 1), character.shape[0] // (character_count + 1)), interpolation = cv2.INTER_AREA)
    #     cv2.imwrite(f'character/{i + man_count + 1}.png', character)
    #     os.remove('out.png')
    #     wait_time = time.time() - start_time
    #     if wait_time < 60:
    #         time.sleep(60 - wait_time)
    
    character_prompt = (',').join(man) + ',' + (',').join(woman)

    first_prompt = f'write the first part of story about {content_text}. write as much as possible. this is about these people\'s story. \"{character_prompt}\"'
    general_prompt = f'write the next part as much as possible. this is about these people\'s story. \"{character_prompt}\"'
    last_prompt = f'write the next part as much as possible. this is about these people\'s story. \"{character_prompt}\"'
    model = musicgen.MusicGen.get_pretrained('medium', device='cuda')

    # generate screenplays
    
    chapter_count = st.session_state.chapter_count
    file_list = []
    if os.path.exists('out'):
        shutil.rmtree('out')
    os.mkdir('out')
    for i in range(chapter_count):
        prompt = ''
        if i == 0:  
            prompt = first_prompt
            response = client.chat.completions.create(
                model="gpt-4-1106-preview",
                messages=[
                    {"role": "user", "content": prompt},
                ]
            )
        elif i == chapter_count - 1:    
            prompt = last_prompt
            last_text = st.session_state.screenplays[len(st.session_state.screenplays) - 1]

            response = client.chat.completions.create(
                model="gpt-4-1106-preview",
                messages=[
                    {"role": "assistant", "content": last_text},
                    {"role": "user", "content": prompt}
                ]
            )
        else:
            prompt = general_prompt
            last_text = st.session_state.screenplays[len(st.session_state.screenplays) - 1]

            response = client.chat.completions.create(
                model="gpt-4-1106-preview",
                messages=[
                    {"role": "assistant", "content": last_text},
                    {"role": "user", "content": prompt}
                ]
            )

            response.stream_to_file("output.mp3")

        print (response.choices[0].message.content)
        
        screenplay = response.choices[0].message.content
        st.session_state.screenplays.append(screenplay)

        cartoon_prompts = screenplay.split('\n\n')
        screenplay_expander = st.expander(f'Screenplay Part {i + 1}')
        screenplay_expander.text(screenplay)
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')  # Specify the video codec
        video = cv2.VideoWriter(f'out.mp4', fourcc, 24, (1024, 1024))
        audio = AudioSegment.silent(0)
        back_audio = []
        for (j, cartoon_prompt) in enumerate(cartoon_prompts):
            start_time = time.time()
            response = client.images.generate(
                model="dall-e-2",
                prompt=f'cartoon image about \'{cartoon_prompt}\'',
                size="1024x1024",
                quality="standard",
                n=1,
                )
            image_data = requests.get(response.data[0].url).content
            screenplay_expander.image(image_data)
            response = client.audio.speech.create(
                model="tts-1",
                voice="nova",
                input=cartoon_prompt,
            )
            response.stream_to_file('out.mp3')
            audio_bytes = open('out.mp3', 'rb').read()
            screenplay_expander.audio(audio_bytes, format='audio/mp3')
            speaker_data = AudioSegment.from_file('out.mp3')
            audio += speaker_data
            model.set_generation_params(duration=len(speaker_data)/1000.0)
            back_res = model.generate([cartoon_prompt], progress=True)
            wavfile.write('out.wav', 32000, back_res.cpu().numpy())
            temp_audio = AudioFileClip('out.wav')
            back_audio.append(temp_audio)

            if 12 - time.time() + start_time > 0 :
                time.sleep(12 - time.time() + start_time)
            
            image_data = Image.open(BytesIO(image_data))
            image_data.save('back.png')
            character_images = []
            for i in range(character_count):
                character_images.append(f'character/{i + 1}.png')
            background = imageCompose('back.png', character_images)

            image = cv2.cvtColor(np.array(background), cv2.COLOR_RGB2BGR)
            duration = int(24 * len(speaker_data) / 1000)
            for _ in range(duration):
                video.write(image)
            if j == 1:
                break
        video.release()
        audio.export('out.mp3', format='mp3')

        video_with_audio = VideoFileClip('out.mp4')
        audio_data = AudioFileClip('out.mp3')
        video_with_audio = video_with_audio.set_audio(audio_data)
        video_with_audio.write_videofile(f'out/{i + 1}.mp4', codec="libx264", audio_codec="aac")
        file_list.append(f'out/{i + 1}.mp4')
        screenplay_document.add_paragraph(screenplay)
    
    video_clips = [VideoFileClip(file) for file in file_list]
    final_clip = concatenate_videoclips(video_clips)
    original_audio = final_clip.audio
    
    combined_audio = concatenate_audioclips(back_audio)
    combined_audio = combined_audio.volumex(0.15)
    if combined_audio.duration > final_clip.duration:
        combined_audio = combined_audio.subclip(0, final_clip.duration)
    fianl_audio = CompositeAudioClip([original_audio, combined_audio])
    final_clip = final_clip.set_audio(fianl_audio)

    final_clip.write_videofile('out_temp.mp4', codec="libx264")

    screenplay_document.save('screenplay.docx')
    
    st.session_state.file_available = True
    st.experimental_rerun()

if st.session_state.file_available:
    for i in range(st.session_state.chapter_count):
        screenplay_expander = st.expander(f'Screenplay Part {i + 1}')
        screenplay_expander.text(st.session_state.screenplays[i])

        # expander.image(st.session_state.images[i])
    col1, col2, col3, col4, col5 = st.columns(5)
    with open("screenplay.docx", "rb") as file:
        col2.download_button(
                label="Download screenplay",
                data=file,
                file_name="screenplay.docx"
            )
    with open("out.mp4", "rb") as file:
        col4.download_button(
                label="Download video",
                data=file,
                file_name="out.mp4"
            )