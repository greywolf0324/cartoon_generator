import streamlit as st
import openai
from apikey import *
import os
import requests
from docx import Document
from docx.shared import Inches
from io import BytesIO
import docx2pdf
from PIL import Image

import torch
from diffusers import StableDiffusionPipeline


if os.name == 'nt':
    import pythoncom
    pythoncom.CoInitialize()
if os.name == 'posix':
    import subprocess
    def convert_docx_to_pdf(docx_path, pdf_path):
        subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_path, docx_path])

openai.api_key = decrypt(OPENAI_API_KEY)

def get_image_data(url):
    response = requests.get(url)
    response.raise_for_status()
    return response.content

if 'book' not in st.session_state:
    st.session_state.book = []
if 'chapter_count' not in st.session_state:
    st.session_state.chapter_count = 1
if 'file_available' not in st.session_state:
    st.session_state.file_available = False
if 'paragraphs' not in st.session_state:
    st.session_state.paragraphs = []
if 'images' not in st.session_state:
    st.session_state.images = []
if 'pipe' not in st.session_state:
    if torch.cuda.is_available():
        device = 'cuda'
    else:
        device = 'cpu'
    pipe = StableDiffusionPipeline.from_pretrained("runwayml/stable-diffusion-v1-5", user_auth_token=True)
    st.session_state.pipe = pipe.to(device)

# initialize document format
document = Document()
section = document.sections[0]
section.page_width = Inches(8.27)  # Width for A4 paper
section.page_height = Inches(11.69)  # Height for A4 paper

section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)


st.write('Book Title')
content_text = st.text_input('Write your contents of your book')
chapter_count = st.text_input('Write the count of chapters')
generate_btn = st.button('Generate books')

if content_text and generate_btn and chapter_count:
    st.session_state.file_available = False
    st.session_state.chapter_count = int(chapter_count)
    first_prompt = f'writhe the first chapter of the book about {content_text}. write as much as possible. start the text with chapter 1'
    general_prompt = 'write the only next chapter and write as much as possible'
    last_prompt = 'write the only next chapter and end the book'
    
    # generate paragraph
    chapter_count = st.session_state.chapter_count
    for i in range(chapter_count):
        prompt = ''
        if i == 0:  
            prompt = first_prompt
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": prompt},
                ]
            )
        elif i == chapter_count - 1:    
            prompt = last_prompt
            last_text = st.session_state.book[len(st.session_state.book) - 1]

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "assistant", "content": last_text},
                    {"role": "user", "content": prompt}
                ]
            )
        else:
            prompt = general_prompt
            last_text = st.session_state.book[len(st.session_state.book) - 1]

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "assistant", "content": last_text},
                    {"role": "user", "content": prompt}
                ]
            )

        # generate the brief case of sentence
        paragraph = response['choices'][0]['message']['content']
        st.session_state.book.append(paragraph)
        expander = st.expander(f'Chapter {i + 1}')
        expander.text(paragraph)
        response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "user", "content": f'\"{paragraph}\" I want to generate image from this paragraph. Give me one sentence prompt that represent for this paragraph'}
                ]
            )
        first_sentence = response['choices'][0]['message']['content']

        # Generate Images

        # response = openai.Image.create(
        #     prompt=first_sentence,
        #     n=1,
        #     size="512x512"
        #     )
        # image_url = response['data'][0]['url']
        # image_data = get_image_data(image_url)

        # image_stream = BytesIO(image_data)
        # img = Image.open(image_stream)
        img = st.session_state.pipe(first_sentence).images[0]
        crop_image = img.crop((0, 0, 512, 384))

        img_stream = BytesIO()
        crop_image.save(img_stream, format='JPEG')
        document.add_picture(img_stream, Inches(6.27))
        expander.image(img_stream)
        st.session_state.images.append(img_stream)

        document.add_paragraph(paragraph)
        st.session_state.paragraphs.append(paragraph)

    document.save('output.docx')
    if os.name == 'nt':
        docx2pdf.convert('output.docx', 'output.pdf')
    if os.name == 'posix':
        convert_docx_to_pdf('output.docx', 'output.pdf')
    st.session_state.file_available = True
    st.experimental_rerun()

if st.session_state.file_available:
    for i in range(st.session_state.chapter_count):
        expander = st.expander(f'Chapter {i + 1}')
        expander.text(st.session_state.paragraphs[i])
        expander.image(st.session_state.images[i])
    
    with open("output.docx", "rb") as file:
        btn = st.download_button(
                label="Download docx",
                data=file,
                file_name="output.docx"
            )
    with open("output.pdf", "rb") as file:
        btn = st.download_button(
                label="Download pdf",
                data=file,
                file_name="output.pdf"
            )